from flask import Flask, render_template, send_file, request, redirect, url_for
import pandas as pd
import os
import matplotlib.pyplot as plt
from statsmodels.tsa.arima.model import ARIMA
from statsmodels.tsa.stattools import adfuller
from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
from statsmodels.tsa.seasonal import seasonal_decompose
from docx import Document
from docx.shared import Inches
import glob

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'csv'}

# Pastikan folder untuk upload ada
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Fungsi untuk memeriksa ekstensi file
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Fungsi preprocessing: penggabungan dan pembersihan data
def preprocess_data(file_path):
    # Membaca file CSV yang diupload
    data = pd.read_csv(file_path, encoding='ISO-8859-1', on_bad_lines='skip', sep=None, engine='python')

    # Menambahkan penggabungan dan pembersihan data yang sudah ada sebelumnya
    data['Tanggal'] = pd.to_datetime(data['Tanggal'], errors='coerce')
    diagnosa_columns = ['Diagnosa 1']
    for col in diagnosa_columns:
        data[col] = data[col].fillna('')

    data['All_Diagnoses'] = data[diagnosa_columns].apply(lambda x: ' | '.join(x), axis=1)
    data['All_Diagnoses'] = data['All_Diagnoses'].str.replace(' \| ', ' | ')

    # Menghapus duplikat berdasarkan kolom tertentu
    data.drop_duplicates(subset=['No. eRM', 'Tanggal'], keep='first', inplace=True)

    # Menggabungkan data untuk analisis lebih lanjut
    data_selected = data[['Tanggal', 'All_Diagnoses']].copy()
    data_selected.rename(columns={'All_Diagnoses': 'Diagnosa 1'}, inplace=True)
    data_selected['Tanggal'] = pd.to_datetime(data_selected['Tanggal'], errors='coerce').dt.date

    # Menghapus baris kosong atau yang hanya berisi spasi
    data_cleaned = data_selected.dropna(subset=['Tanggal', 'Diagnosa 1'])
    data_cleaned = data_cleaned[data_cleaned['Diagnosa 1'].str.strip() != '']

    # Agregasi data berdasarkan Tanggal dan Diagnosa
    data_aggregated = data_cleaned.groupby(['Tanggal', 'Diagnosa 1']).size().reset_index(name='Total Kasus')

    # Menyimpan data yang telah dibersihkan dan diagregasi ke file CSV baru
    output_file_clean = os.path.join(app.config['UPLOAD_FOLDER'], 'data_aggregated.csv')
    data_aggregated.to_csv(output_file_clean, index=False)

    return output_file_clean

# Fungsi untuk analisis dan pembuatan dokumen
def generate_report(file_path):
    # Load data dari file yang telah dibersihkan dan diagregasi
    data = pd.read_csv(file_path, parse_dates=['Tanggal'])

    # Hitung jumlah total kasus per diagnosa
    diagnosa_counts = data.groupby('Diagnosa 1')['Total Kasus'].sum().reset_index()

    # Ambil 10 diagnosa teratas berdasarkan jumlah kasus
    top_10_diagnoses = diagnosa_counts.sort_values(by='Total Kasus', ascending=False).head(10)

    # File output untuk Word
    output_file = "Top10_Diagnoses_Analysis_Weekly.docx"
    doc = Document()
    doc.add_heading("Analisis 10 Diagnosis Teratas (Resampling Mingguan)", level=1)

    # Proses untuk setiap diagnosis
    for diagnosis_terpilih in top_10_diagnoses['Diagnosa 1']:
        doc.add_heading(f"Diagnosis: {diagnosis_terpilih}", level=2)

        # Filter data berdasarkan diagnosis
        data_diagnosis = data[data['Diagnosa 1'] == diagnosis_terpilih]
        data_diagnosis.set_index('Tanggal', inplace=True)

        # Resample data menjadi mingguan
        weekly_data = data_diagnosis['Total Kasus'].resample('W').sum().fillna(0)
        doc.add_heading("Data Mingguan", level=3)
        doc.add_paragraph(weekly_data.to_string())

        # Uji ADF (Stationarity Test)
        adf_test = adfuller(weekly_data)
        adf_result = {
            "Test Statistic": adf_test[0],
            "p-Value": adf_test[1],
            "Lags Used": adf_test[2],
            "Number of Observations Used": adf_test[3]
        }
        doc.add_heading("Hasil Uji ADF", level=3)
        for key, value in adf_result.items():
            doc.add_paragraph(f"{key}: {value}")

        # Jika data tidak stationer, lakukan differencing
        if adf_test[1] > 0.05:
            weekly_data_diff = weekly_data.diff().dropna()
            doc.add_paragraph("Data tidak stationer. Differencing dilakukan.")
        else:
            weekly_data_diff = weekly_data

        # Grafik dekomposisi musiman
        try:
            decomposition = seasonal_decompose(weekly_data, model='additive', period=26)  # 52 minggu dalam setahun
            decomposition.plot()
            seasonal_plot = f"{diagnosis_terpilih[:20]}_seasonal_weekly.png"
            plt.savefig(seasonal_plot)
            plt.close()

            # Tambahkan grafik ke dokumen
            doc.add_heading("Grafik Dekomposisi Musiman", level=3)
            doc.add_picture(seasonal_plot, width=Inches(6))
            os.remove(seasonal_plot)
        except Exception as e:
            doc.add_paragraph("Dekomposisi musiman gagal dilakukan.")
            doc.add_paragraph(f"Error: {e}")

        # Plot ACF dan PACF
        plt.figure(figsize=(12, 6))
        plot_acf(weekly_data_diff, ax=plt.subplot(121), title='ACF')
        plot_pacf(weekly_data_diff, ax=plt.subplot(122), title='PACF')
        plt.suptitle(f"ACF dan PACF untuk {diagnosis_terpilih} (Mingguan)", fontsize=16)
        plt.tight_layout()
        acf_pacf_plot = f"{diagnosis_terpilih[:20]}_acf_pacf_weekly.png"
        plt.savefig(acf_pacf_plot)
        plt.close()

        # Tambahkan gambar ACF dan PACF ke dokumen
        doc.add_heading("Plot ACF dan PACF", level=3)
        doc.add_picture(acf_pacf_plot, width=Inches(6))
        os.remove(acf_pacf_plot)

        # Model ARIMA
        try:
            model = ARIMA(weekly_data, order=(1, 1, 1))
            model_fit = model.fit()
            mse = ((model_fit.resid) ** 2).mean()
            doc.add_heading("Model ARIMA", level=3)
            doc.add_paragraph(f"Mean Squared Error (MSE): {mse}")

            # Prediksi 12 minggu ke depan
            forecast = model_fit.forecast(steps=12)
            doc.add_heading("Prediksi 12 Minggu ke Depan", level=3)
            doc.add_paragraph(forecast.to_string())
        except Exception as e:
            doc.add_paragraph(f"Model ARIMA gagal untuk {diagnosis_terpilih}. Error: {e}")

    # Simpan dokumen
    doc.save(output_file)
    return output_file

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return redirect(request.url)
    
    file = request.files['file']
    if file.filename == '':
        return redirect(request.url)
    
    if file and allowed_file(file.filename):
        # Simpan file yang di-upload
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)

        # Preprocess data (penggabungan dan pembersihan)
        cleaned_file_path = preprocess_data(file_path)

        # Jalankan analisis dan buat dokumen
        output_file = generate_report(cleaned_file_path)
        return render_template('index.html', file_path=output_file)

    return "File tidak valid, pastikan file CSV yang di-upload benar."

@app.route('/download')
def download():
    # Menyediakan file untuk diunduh
    return send_file('Top10_Diagnoses_Analysis_Weekly.docx', as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
